"""
AI-powered applicant ranking service using Google Gemini
"""
import logging
import json
import io
import tempfile
import os
from typing import List, Dict, Optional
from pathlib import Path
import google.generativeai as genai
from app.config import settings
from app.models import Application
from app.storage import gcp_storage

# Document processing imports
try:
    from PyPDF2 import PdfReader
    from docx import Document
    from PIL import Image
    from pdf2image import convert_from_path
    import pytesseract
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError as e:
    logger = logging.getLogger(__name__)
    logger.warning(f"Document processing libraries not available: {e}")
    DOCUMENT_PROCESSING_AVAILABLE = False

logger = logging.getLogger(__name__)


class ResumeExtractor:
    """Extract text content from various resume formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            text = ""
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # If no text extracted, try OCR
            if not text.strip():
                logger.info("No text in PDF, attempting OCR")
                text = ResumeExtractor.extract_text_with_ocr(file_content)
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_with_ocr(file_content: bytes) -> str:
        """Extract text from PDF using OCR"""
        try:
            # Save to temporary file for pdf2image
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                # Convert PDF to images
                images = convert_from_path(tmp_path)
                text = ""
                
                for i, image in enumerate(images):
                    logger.info(f"Performing OCR on page {i+1}")
                    page_text = pytesseract.image_to_string(image)
                    text += page_text + "\n"
                
                return text.strip()
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                    
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_resume_content(resume_blob_path: str) -> Optional[str]:
        """Download and extract text from resume"""
        try:
            # Download file from GCS
            if not gcp_storage.client or not gcp_storage.bucket:
                logger.error("GCS not initialized")
                return None
            
            blob = gcp_storage.bucket.blob(resume_blob_path)
            file_content = blob.download_as_bytes()
            
            # Determine file type and extract accordingly
            file_ext = Path(resume_blob_path).suffix.lower()
            
            if file_ext == '.pdf':
                return ResumeExtractor.extract_text_from_pdf(file_content)
            elif file_ext in ['.docx', '.doc']:
                return ResumeExtractor.extract_text_from_docx(file_content)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to extract resume content: {e}")
            return None


class ApplicantAIAnalyzer:
    """Analyzes and ranks job applicants using Gemini AI"""
    
    def __init__(self):
        self.api_key = settings.GEMINI_API_KEY
        if self.api_key:
            try:
                genai.configure(api_key=self.api_key)
                self.model = genai.GenerativeModel('gemini-2.0-flash-exp')
                logger.info("Gemini AI initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize Gemini AI: {e}")
                self.model = None
        else:
            logger.warning("GEMINI_API_KEY not configured")
            self.model = None
    
    def analyze_applicants(
        self, 
        applications: List[Application], 
        job_title: str,
        job_description: str,
        job_requirements: Optional[str] = None
    ) -> List[Dict]:
        """
        Analyze and rank applicants using AI with resume content extraction
        
        Returns list of applications with AI scores and insights
        """
        if not self.model:
            logger.warning("Gemini AI not available, returning unranked applications")
            return self._fallback_ranking(applications)
        
        if not applications:
            return []
        
        try:
            # Prepare applicant data with resume content extraction
            applicants_data = []
            for app in applications:
                resume_content = None
                if app.resume_url and DOCUMENT_PROCESSING_AVAILABLE:
                    logger.info(f"Extracting resume for application {app.id}")
                    resume_content = ResumeExtractor.extract_resume_content(app.resume_url)
                    if resume_content:
                        logger.info(f"Successfully extracted {len(resume_content)} characters from resume")
                    else:
                        logger.warning(f"Failed to extract resume content for application {app.id}")
                
                applicant_info = {
                    "id": app.id,
                    "name": f"{app.applicant.first_name} {app.applicant.last_name or ''}".strip(),
                    "username": app.applicant.username or "N/A",
                    "cover_letter": app.cover_letter or "No cover letter provided",
                    "has_resume": bool(app.resume_url),
                    "resume_content": resume_content if resume_content else "Resume not accessible or no text content",
                    "status": app.status.value
                }
                applicants_data.append(applicant_info)
            
            # Create comprehensive prompt for Gemini with resume analysis
            prompt = self._create_analysis_prompt(
                job_title=job_title,
                job_description=job_description,
                job_requirements=job_requirements,
                applicants=applicants_data
            )
            
            # Get AI analysis
            logger.info("Sending request to Gemini AI for applicant analysis")
            response = self.model.generate_content(prompt)
            analysis_text = response.text
            
            # Parse the AI response
            ranked_applicants = self._parse_ai_response(analysis_text, applications)
            
            logger.info(f"Successfully analyzed {len(ranked_applicants)} applicants")
            return ranked_applicants
            
        except Exception as e:
            logger.error(f"AI analysis failed: {e}")
            return self._fallback_ranking(applications)
    
    def _create_analysis_prompt(
        self, 
        job_title: str, 
        job_description: str,
        job_requirements: Optional[str],
        applicants: List[Dict]
    ) -> str:
        """Create a detailed prompt for Gemini to analyze applicants with resume content"""
        
        requirements_section = f"\n**Requirements:**\n{job_requirements}" if job_requirements else ""
        
        prompt = f"""You are an expert HR professional and talent acquisition specialist with extensive experience in technical recruiting. Analyze the following job applicants and rank them based on their suitability for the position.

**Job Position:** {job_title}

**Job Description:**
{job_description}
{requirements_section}

**Applicants to Analyze:**
"""
        
        for i, applicant in enumerate(applicants, 1):
            resume_section = f"""
   - **Resume Content:** 
     {applicant['resume_content'][:3000] if len(applicant['resume_content']) > 3000 else applicant['resume_content']}
     {"... (truncated for length)" if len(applicant['resume_content']) > 3000 else ""}
""" if applicant['resume_content'] != "Resume not accessible or no text content" else "\n   - **Resume:** Not available\n"

            prompt += f"""
{i}. **{applicant['name']}** (@{applicant['username']})
   - Application ID: {applicant['id']}
   - Resume Attached: {"Yes" if applicant['has_resume'] else "No"}
   - Cover Letter: "{applicant['cover_letter']}"
{resume_section}
   - Current Status: {applicant['status']}

"""
        
        prompt += """
**Your Task:**
Analyze each applicant comprehensively and provide a JSON response with the following structure for each applicant:

```json
[
  {
    "application_id": <application_id>,
    "overall_score": <0-100>,
    "cover_letter_score": <0-100>,
    "completeness_score": <0-100>,
    "relevance_score": <0-100>,
    "resume_score": <0-100>,
    "ai_summary": "<2-3 sentence summary of why this candidate is a good or poor fit>",
    "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
    "concerns": ["<concern 1>", "<concern 2>"],
    "recommendation": "<hire|interview|maybe|pass>"
  }
]
```

**Scoring Criteria:**
- **overall_score**: Holistic assessment combining all factors (0-100)
  - Give HEAVY WEIGHT to resume content if available (40% of overall score)
  - Cover letter quality (20%)
  - Completeness (20%)
  - Relevance to job requirements (20%)

- **cover_letter_score**: Quality, relevance, and professionalism of cover letter (0-100)
  - Clear communication
  - Genuine interest
  - Relevant examples
  - Professional tone

- **completeness_score**: How complete is the application (0-100)
  - Resume attached (+50 points)
  - Detailed cover letter (+50 points)

- **relevance_score**: How well the applicant's experience/skills match the job requirements (0-100)
  - Based on job description match
  - Required skills present
  - Experience level appropriate

- **resume_score**: Quality and relevance of resume content (0-100)
  - ONLY if resume content is available
  - Technical skills match
  - Relevant experience
  - Education background
  - Projects and achievements
  - Overall professionalism
  - If no resume: score = 0

**Important Instructions for Resume Analysis:**
- If resume content is available, analyze it DEEPLY for:
  - Technical skills and technologies mentioned
  - Years of experience
  - Relevant projects
  - Education and certifications
  - Career progression
  - Specific achievements
- Applicants with strong resumes should score significantly higher
- Mention specific resume details in strengths
- The resume_score should heavily influence the overall_score

**Recommendations:**
- "hire" - Strong candidate with excellent resume and fit, highly recommended (overall_score >= 85)
- "interview" - Good candidate worth interviewing (overall_score >= 70)
- "maybe" - Moderate fit, consider as backup (overall_score >= 50)
- "pass" - Not a good fit for this position (overall_score < 50)

Please provide ONLY the JSON array, no additional text or markdown formatting. Order by overall_score descending (best candidates first).
"""
        
        return prompt
    
    def _parse_ai_response(self, response_text: str, applications: List[Application]) -> List[Dict]:
        """Parse AI response and merge with application data"""
        
        try:
            # Clean the response text
            cleaned_text = response_text.strip()
            
            # Remove markdown code blocks if present
            if cleaned_text.startswith("```json"):
                cleaned_text = cleaned_text[7:]
            elif cleaned_text.startswith("```"):
                cleaned_text = cleaned_text[3:]
            
            if cleaned_text.endswith("```"):
                cleaned_text = cleaned_text[:-3]
            
            cleaned_text = cleaned_text.strip()
            
            # Parse JSON
            ai_rankings = json.loads(cleaned_text)
            
            # Create a map of application_id to application object
            app_map = {app.id: app for app in applications}
            
            # Merge AI data with application data
            ranked_results = []
            for ranking in ai_rankings:
                app_id = ranking.get("application_id")
                if app_id in app_map:
                    app = app_map[app_id]
                    result = {
                        "application": app,
                        "ai_analysis": {
                            "overall_score": ranking.get("overall_score", 0),
                            "cover_letter_score": ranking.get("cover_letter_score", 0),
                            "completeness_score": ranking.get("completeness_score", 0),
                            "relevance_score": ranking.get("relevance_score", 0),
                            "resume_score": ranking.get("resume_score", 0),
                            "ai_summary": ranking.get("ai_summary", ""),
                            "strengths": ranking.get("strengths", []),
                            "concerns": ranking.get("concerns", []),
                            "recommendation": ranking.get("recommendation", "maybe")
                        }
                    }
                    ranked_results.append(result)
            
            return ranked_results
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"Response text: {response_text[:500]}")
            return self._fallback_ranking(applications)
        except Exception as e:
            logger.error(f"Error parsing AI response: {e}")
            return self._fallback_ranking(applications)
    
    def _fallback_ranking(self, applications: List[Application]) -> List[Dict]:
        """Fallback ranking when AI is unavailable - simple rule-based scoring"""
        
        results = []
        for app in applications:
            # Simple scoring based on completeness
            completeness_score = 0
            if app.cover_letter:
                completeness_score += 50
            if app.resume_url:
                completeness_score += 50
            
            # Basic cover letter quality check
            cover_letter_score = 0
            if app.cover_letter:
                length = len(app.cover_letter)
                if length > 200:
                    cover_letter_score = 80
                elif length > 100:
                    cover_letter_score = 60
                elif length > 50:
                    cover_letter_score = 40
                else:
                    cover_letter_score = 20
            
            # Resume score
            resume_score = 50 if app.resume_url else 0
            
            overall_score = int((completeness_score + cover_letter_score + resume_score) / 3)
            
            result = {
                "application": app,
                "ai_analysis": {
                    "overall_score": overall_score,
                    "cover_letter_score": cover_letter_score,
                    "completeness_score": completeness_score,
                    "relevance_score": 50,  # Default neutral score
                    "resume_score": resume_score,
                    "ai_summary": "AI analysis unavailable. Basic scoring applied based on application completeness.",
                    "strengths": ["Application submitted" if completeness_score > 0 else ""],
                    "concerns": ["AI analysis not available"],
                    "recommendation": "review"
                }
            }
            results.append(result)
        
        # Sort by overall score descending
        results.sort(key=lambda x: x["ai_analysis"]["overall_score"], reverse=True)
        
        return results


# Singleton instance
ai_analyzer = ApplicantAIAnalyzer()
